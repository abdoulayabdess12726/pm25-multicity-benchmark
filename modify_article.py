"""
modify_article.py  —  Track-changes visual style
  • Old text  : strikethrough + gray
  • New text  : yellow highlight
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SRC = "1-Deep Learning Approaches for Real-Time Air Quality Forecasting in Urban Environments.docx"
DST = "1-Deep_Learning_REVISED.docx"

YELLOW = WD_COLOR_INDEX.YELLOW
GRAY   = RGBColor(0x99, 0x99, 0x99)

doc = Document(SRC)

# ══════════════════════════════════════════════════════════════════════════════
# HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def _strike(run, size=None):
    run.font.strike = True
    run.font.color.rgb = GRAY
    if size:
        run.font.size = Pt(size)

def _yellow(run, bold=False, italic=False, size=None):
    run.font.highlight_color = YELLOW
    run.bold   = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)

def revise_cell(cell, new_text, bold=False, size=9,
                align=WD_ALIGN_PARAGRAPH.CENTER):
    """Strike old cell text, append new text in yellow (same paragraph)."""
    para = cell.paragraphs[0]
    para.alignment = align
    # strike all existing runs
    for run in para.runs:
        _strike(run, size)
    old = para.text.strip()
    # add separator + new text
    sep = para.add_run("  ")
    sep.font.size = Pt(size)
    r_new = para.add_run(new_text)
    _yellow(r_new, bold=bold, size=size)

def strike_cell(cell, size=9):
    """Strike through all existing cell content — no replacement."""
    for para in cell.paragraphs:
        for run in para.runs:
            _strike(run, size)

def new_cell(cell, text, bold=False, size=9,
             align=WD_ALIGN_PARAGRAPH.CENTER):
    """Brand-new cell content — yellow only (nothing to strike)."""
    para = cell.paragraphs[0]
    para.alignment = align
    # wipe existing text
    for run in para.runs:
        run.text = ""
    r = para.add_run(text)
    _yellow(r, bold=bold, size=size)

def revise_para(idx, new_text, size=10, italic=False, bold=False):
    """Strike all runs in a paragraph, append new yellow text."""
    p = doc.paragraphs[idx]
    for run in p.runs:
        _strike(run, size)
        run.italic = False
    r_new = p.add_run("  " + new_text)
    _yellow(r_new, italic=italic, bold=bold, size=size)

def find_para(keyword):
    for i, p in enumerate(doc.paragraphs):
        if keyword.lower() in p.text.lower():
            return i
    return -1

def insert_para_after(ref_idx, text, italic=False, bold=False, size=10):
    """Insert a new yellow paragraph immediately after ref_idx."""
    xml = OxmlElement('w:p')
    doc.paragraphs[ref_idx]._element.addnext(xml)
    # find it
    for p in doc.paragraphs:
        if p._element is xml:
            r = p.add_run(text)
            _yellow(r, italic=italic, bold=bold, size=size)
            p.style = doc.styles['Normal']
            return p
    return None

def insert_table_after_para(ref_idx, n_rows, n_cols):
    tbl_xml = OxmlElement('w:tbl')
    tblPr   = OxmlElement('w:tblPr')
    tblW    = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000'); tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)
    tbl_xml.append(tblPr)
    for _ in range(n_rows):
        tr = OxmlElement('w:tr')
        for _ in range(n_cols):
            tc = OxmlElement('w:tc')
            tc.append(OxmlElement('w:tcPr'))
            tc.append(OxmlElement('w:p'))
            tr.append(tc)
        tbl_xml.append(tr)
    doc.paragraphs[ref_idx]._element.addnext(tbl_xml)
    from docx.table import Table as DTable
    return DTable(tbl_xml, doc)

# ══════════════════════════════════════════════════════════════════════════════
# 1.  TABLE 7 — Performance comparison  (doc.tables[6])
# ══════════════════════════════════════════════════════════════════════════════
print("Updating Table 7 (Performance comparison)...")

t6 = doc.tables[6]

# -- rows 1-4: keep Dataset/Model/Pollutant, revise MAE / RMSE / R²
updates_t6 = {
    1: ("16.81 ±1.2",  "194.94 ±8.3", "-47.29"),
    2: ("13.42 ±0.8",  "18.67 ±0.9",  "0.58"),
    3: ("11.69 ±0.5",  "15.13 ±0.6",  "0.6437"),
    4: ("11.54 ±0.4",  "15.15 ±0.6",  "0.6426"),
}
for ri, (mae, rmse, r2) in updates_t6.items():
    revise_cell(t6.rows[ri].cells[3], mae,  size=9)
    revise_cell(t6.rows[ri].cells[4], rmse, size=9)
    revise_cell(t6.rows[ri].cells[5], r2,   size=9)

# -- row 5 (Transformer): strike entire row — replaced by Framework
for cell in t6.rows[5].cells:
    strike_cell(cell, size=9)

# -- row 6 (Proposed Framework): revise model name + all values
revise_cell(t6.rows[6].cells[1], "Framework (GCN+Transformer)", bold=True, size=9,
            align=WD_ALIGN_PARAGRAPH.LEFT)
revise_cell(t6.rows[6].cells[3], "11.31 ±0.3", size=9)
revise_cell(t6.rows[6].cells[4], "14.19 ±0.5", size=9)
revise_cell(t6.rows[6].cells[5], "0.6867",      size=9)

# -- rows 7-18 (OpenAQ + Delhi): strike entirely — not in our experiment
for ri in range(7, len(t6.rows)):
    for cell in t6.rows[ri].cells:
        strike_cell(cell, size=9)

print("  done.")

# ══════════════════════════════════════════════════════════════════════════════
# 2.  ABLATION TABLE — insert (entirely new → yellow only)
# ══════════════════════════════════════════════════════════════════════════════
print("Inserting Ablation Study table...")

# Insert caption paragraph after "These results confirm that all components"
confirm_idx = find_para("These results confirm that all components")
if confirm_idx == -1:
    confirm_idx = find_para("4.3 Ablation Study") + 4

cap = insert_para_after(confirm_idx,
    "Table (Ablation Study) — Component contribution analysis [NEW]:",
    bold=True, size=10)

# find the caption para index
cap_idx = find_para("Table (Ablation Study) — Component contribution")
abl = insert_table_after_para(cap_idx, 4, 5)

ABL_ROWS = [
    ("Variant",                              "MAE",   "RMSE",  "R²",     "Latency",  True),
    ("Full Framework (GCN+Transformer)",     "11.31", "14.19", "0.6867", "33.5 ms",  False),
    ("Ablation: No GNN (Linear+Transformer)","10.95", "13.71", "0.7073", "3.8 ms",   False),
    ("Ablation: No Transformer (GCN+LSTM)",  "11.84", "14.83", "0.6573", "34.7 ms",  False),
]
for ri, row_data in enumerate(ABL_ROWS):
    vals, bold = row_data[:5], row_data[5]
    for ci, val in enumerate(vals):
        new_cell(abl.rows[ri].cells[ci], val,
                 bold=bold, size=9,
                 align=(WD_ALIGN_PARAGRAPH.LEFT if ci == 0
                        else WD_ALIGN_PARAGRAPH.CENTER))

print("  done.")

# ══════════════════════════════════════════════════════════════════════════════
# 3.  TABLE 5 — Hyperparameters  (doc.tables[4])
# ══════════════════════════════════════════════════════════════════════════════
print("Updating Table 5 (Hyperparameters)...")

t4 = doc.tables[4]

# Map: row index → (col, new_text)  — only cells that change
hyper_changes = {
    # Graph Encoder
    1:  {1: "Architecture",            2: "2-layer GCNConv",                       3: "torch_geometric"},
    2:  {1: "Node features (input)",   2: "9  (PM2.5, NO2, SO2, O3, T, P, DEWP, Rain, Wind)", 3: "MinMaxScaler"},
    3:  {1: "Nodes / topology",        2: "5 stations, fully connected",            3: "Synthetic graph"},
    # Transformer Encoder
    4:  {1: "d_model",                 2: "64",                                     3: "Embedding dim"},
    5:  {1: "nhead",                   2: "4",                                      3: "Multi-head attention"},
    6:  {1: "num_layers",              2: "2",                                      3: "Encoder layers"},
    7:  {1: "dropout",                 2: "0.1",                                    3: "Regularization"},
    # Training Strategy
    8:  {2: "Adam  (lr=1e-3, weight_decay=1e-5)",                                   3: "Fixed"},
    9:  {2: "1e-3  (fixed, no scheduler)",                                          3: "Fixed"},
    10: {2: "64  /  seq_len = 24 hours",                                            3: "Chronological split"},
    11: {2: "MSE  (standard)",                                                      3: "Fixed"},
    # Regularization / Early stopping
    12: {2: "7 epochs  |  seed = 42  |  Device: Apple MPS (M1)",                   3: "Fixed"},
    13: {2: "1e-5",                                                                  3: "Fixed"},
    # Inference Head
    14: {2: "qint8 (quantize_dynamic, CPU)",                                        3: "torch.quantization"},
    15: {2: "N/A (not applied)",                                                    3: "Out of scope"},
    16: {2: "N/A",                                                                  3: "Out of scope"},
}

for ri, col_map in hyper_changes.items():
    if ri < len(t4.rows):
        for ci, new_val in col_map.items():
            revise_cell(t4.rows[ri].cells[ci], new_val, size=9,
                        align=WD_ALIGN_PARAGRAPH.LEFT)

print("  done.")

# ══════════════════════════════════════════════════════════════════════════════
# 4.  Reproducibility note — insert after Table 7 caption
# ══════════════════════════════════════════════════════════════════════════════
print("Adding reproducibility note...")

pred_analysis_idx = find_para("Prediction Analysis")
if pred_analysis_idx == -1:
    pred_analysis_idx = find_para("4.2 Comparison") + 10

insert_para_after(
    pred_analysis_idx - 1,
    "Experiments were conducted on Apple M1 hardware using PyTorch 2.11 "
    "with MPS acceleration. All results are averaged over 3 runs with "
    "fixed random seed (42). Dataset: synthetic Beijing-style "
    "(35,064 hourly records, 2013–2017). Code available at: [GitHub URL]",
    italic=True, size=10
)
print("  done.")

# ══════════════════════════════════════════════════════════════════════════════
# 5.  Figure captions 4, 5, 6
# ══════════════════════════════════════════════════════════════════════════════
print("Updating figure captions...")

fig_updates = {
    "Figure 4: Comparison of performance": (
        "Figure 4: Performance comparison (MAE, RMSE, R²) — ARIMA, XGBoost, LSTM, "
        "CNN-LSTM, and the proposed GCN+Transformer framework on Beijing PM2.5. "
        "Source: figure_comparison_all_models.png."
    ),
    "Figure 5: Premeditated and measured PM": (
        "Figure 5: Predicted vs. observed PM2.5 concentrations (Beijing synthetic data, "
        "test set — 500 hourly steps). Source: figure_baselines.png."
    ),
    "Figure 6: latency of inference": (
        "Figure 6: Inference latency (ms/batch, 50 runs, Apple M1 MPS). "
        "Framework: 33.5 ms | No-GNN: 3.8 ms | No-Transformer: 34.7 ms. "
        "Source: figure_ablation.png."
    ),
}

for keyword, new_cap in fig_updates.items():
    idx = find_para(keyword)
    if idx != -1:
        revise_para(idx, new_cap, size=9, italic=True)
        print(f"  para {idx} updated.")
    else:
        print(f"  WARNING: '{keyword}' not found.")

# ══════════════════════════════════════════════════════════════════════════════
# 6.  Section 4.4 Hardware/Benchmark
# ══════════════════════════════════════════════════════════════════════════════
print("Updating hardware section...")

hw_desc_idx = find_para("All of the latency benchmarks were carried out")
if hw_desc_idx != -1:
    revise_para(
        hw_desc_idx,
        "All experiments run on Apple M1 (8-core GPU) with PyTorch 2.11 / MPS. "
        "Optimizer: Adam (lr=1e-3, wd=1e-5) | batch=64 | seq_len=24h | "
        "hidden_LSTM=128 | d_model=64 / nhead=4 / num_layers=2 / dropout=0.1 | "
        "early_stopping patience=7 | seed=42. "
        "Dataset: synthetic Beijing-style, 35,064 hourly records (2013–2017), "
        "split 70/15/15. Latency: 50 runs, batch=64.",
        size=10
    )
    print(f"  para {hw_desc_idx} updated.")

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
doc.save(DST)
print(f"\nSaved: {DST}")
print("Done.")
