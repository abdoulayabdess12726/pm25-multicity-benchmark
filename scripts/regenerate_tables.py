#!/usr/bin/env python3
"""
scripts/regenerate_tables.py — Tables 1 à 11 du manuscrit (sauf T3, dans
scripts/build_hyperparameter_table.py), depuis results/raw_results.csv
EXCLUSIVEMENT pour les tables de résultats (P4, REVISION_BRIEF.md).

Aucun chiffre en dur dans les tables de résultats : chaque valeur affichée
est recalculée depuis raw_results.csv à chaque exécution (agrégats, deltas,
tests statistiques).

EXCEPTIONS DOCUMENTÉES — Table 1 (caractérisation), Table 2 (h(D)) et la
composante h(D) de la Table 7 : ce ne sont PAS des résultats de run (pas de
seed/model/rmse), volontairement exclus du schéma raw_results.csv en P3
(cf. CHANGELOG_TABLES.md). Sources : analysis/p9_3_characterization.csv
(Table 1) et results/heterogeneity_index_v2.csv, 05_compute_heterogeneity_v2.py
(Table 2). Ce sont les deux seules lectures hors raw_results.csv de ce
script — si ce n'est pas ce que "lecture depuis raw_results.csv uniquement"
voulait dire, c'est un point à trancher avant de committer, pas une décision
prise en silence.

Numérotation (renumérotée 2026-08-24, cf. README.md — table de
correspondance complète avec l'ancienne numérotation dans le rapport de
tâche livré à l'utilisateur) :
  T1 caractérisation des réseaux (R2.7) · T2 h(D) · T3 hyperparamètres (R2.3,
  scripts/build_hyperparameter_table.py, PAS ce script) · T4 benchmark ·
  T5 baselines externes · T6 tests statistiques · T7 corrélation inter-villes ·
  T8 ΔR² par station · T9 sensibilité k · T10 over-smoothing/GAT ·
  T11 diagnostics.

RÈGLE D'ARRÊT : si une assertion échoue, on corrige les données ou le
pipeline, JAMAIS l'assertion. Une condition sans aucune ligne dans
raw_results.csv est rapportée MISSING DATA — jamais un succès par vacuité,
jamais ignorée.

Sortie : manuscript/tables/table{1,2,4..11}_*.md + .docx (T3 : voir
scripts/build_hyperparameter_table.py)
Usage : python scripts/regenerate_tables.py
"""
import sys
import warnings
from pathlib import Path

import numpy as np
import pandas as pd
from scipy import stats

warnings.filterwarnings("ignore")

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))
from src.results_io import load_results
from src.stats import agg_mean_std

try:
    from docx import Document
    from docx.shared import Pt
    HAVE_DOCX = True
except ImportError:
    HAVE_DOCX = False

OUT_DIR = ROOT / "manuscript" / "tables"
OUT_DIR.mkdir(parents=True, exist_ok=True)

CITIES = ["beijing", "london", "madrid"]
# T10 (over-smoothing/GAT, E13) et T11 (diagnostics, E4/E5) restent à 3
# réseaux : CZT n'a ni contrôle over-smoothing ni pruning, jamais lancés pour
# ce réseau (E16 : protocole limité à GCN-Transformer + Linear-Transformer).
# CITIES_4NET couvre T2(h(D))/T4(benchmark)/T6(tests stat.)/T7(corrélation
# inter-villes)/T8(par station)/T9(k-sensitivity), où CZT est applicable.
CITIES_4NET = CITIES + ["czt"]
CITY_LABEL = {"beijing": "Beijing", "london": "London", "madrid": "Madrid", "czt": "Chang-Zhu-Tan"}
TOPOS = ["distance", "correlation"]
SEEDS = [42, 123, 777]
KS = [3, 5, 8]
PRIMARY_SEED = 42
TOL_EXACT = 1e-9

H_INDEX_PATH = ROOT / "results" / "heterogeneity_index_v2.csv"

# --------------------------------------------------------------------------- #
# Suivi des assertions — PASS / FAIL / MISSING DATA, jamais un 4e état
# --------------------------------------------------------------------------- #
ASSERTIONS = []


def record(name, status, detail=""):
    assert status in ("PASS", "FAIL", "MISSING DATA"), status
    ASSERTIONS.append(dict(name=name, status=status, detail=detail))
    tag = {"PASS": "OK  ", "FAIL": "FAIL", "MISSING DATA": "MISS"}[status]
    print(f"  [{tag}] {name}" + (f" — {detail}" if detail else ""))


# --------------------------------------------------------------------------- #
# Écriture Markdown + Word (docx natif — ouvrir, sélectionner le tableau,
# copier-coller dans le manuscrit préserve la mise en forme Word).
# --------------------------------------------------------------------------- #
def write_table(name, title, headers, rows, note=""):
    md_path = OUT_DIR / f"{name}.md"
    lines = [f"# {title}\n", "| " + " | ".join(headers) + " |",
             "|" + "|".join(["---"] * len(headers)) + "|"]
    for row in rows:
        lines.append("| " + " | ".join(str(c) for c in row) + " |")
    if note:
        lines.append(f"\n_{note}_")
    md_path.write_text("\n".join(lines) + "\n")

    if HAVE_DOCX:
        doc = Document()
        h = doc.add_heading(title, level=2)
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Light Grid Accent 1"
        for i, htext in enumerate(headers):
            table.rows[0].cells[i].text = str(htext)
        for row in rows:
            cells = table.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = str(val)
        if note:
            p = doc.add_paragraph(note)
            p.runs[0].italic = True
            p.runs[0].font.size = Pt(9)
        doc.save(OUT_DIR / f"{name}.docx")
    print(f"  écrit : {md_path.name}" + (" + .docx" if HAVE_DOCX else " (docx indisponible)"))


def write_missing(name, title, expected_n):
    """Table MISSING DATA explicite — jamais un fichier vide silencieux."""
    md_path = OUT_DIR / f"{name}.md"
    text = (f"# {title}\n\n"
            f"**MISSING DATA** — 0/{expected_n} conditions présentes dans "
            "results/raw_results.csv. Cette table ne peut pas être générée "
            "tant que les runs correspondants n'ont pas été exécutés "
            "(cf. CHANGELOG_TABLES.md, budget de calcul restant).\n")
    md_path.write_text(text)
    if HAVE_DOCX:
        doc = Document()
        doc.add_heading(title, level=2)
        p = doc.add_paragraph(f"MISSING DATA — 0/{expected_n} conditions présentes.")
        p.runs[0].bold = True
        doc.save(OUT_DIR / f"{name}.docx")
    print(f"  écrit (MISSING DATA) : {md_path.name}")


# --------------------------------------------------------------------------- #
# Chargement
# --------------------------------------------------------------------------- #
def load():
    df = load_results()
    if len(df) == 0:
        raise RuntimeError("raw_results.csv vide ou absent — rien à régénérer.")
    for col in ("k", "keep_frac", "n_stations"):
        df[col] = pd.to_numeric(df[col], errors="coerce")
    df["variant"] = df["variant"].fillna("")
    df["topology"] = df["topology"].fillna("")
    df["provenance_note"] = df["provenance_note"].fillna("")
    df = resolve_superseded_suspect_rows(df)
    df = resolve_madrid_correlation_nan_bug_per_station_rows(df)
    return df


_SUSPECT_MARKERS = ("SUSPECT_6STATION", "UNRECOVERABLE")


def _is_suspect(notes):
    """True si au moins une note signale un VRAI problème de provenance
    (SUSPECT_6STATION/UNRECOVERABLE) — pas n'importe quelle note informative
    (ex. "RMSE/MAE non calculés", "source=recompute"). Une première version
    déclenchait [SUSPECT] sur toute note non vide, ce qui flaggait à tort des
    lignes Beijing jamais concernées par MENDEZ ALVARO — corrigé avant E9."""
    return notes.fillna("").str.contains("|".join(_SUSPECT_MARKERS), regex=True).any()


def resolve_superseded_suspect_rows(df):
    """Exclut de la génération des tables les lignes SUSPECT_6STATION/
    UNRECOVERABLE quand un re-run propre existe pour la MÊME condition
    logique (IDENTITY_COLS) sous un run_id différent — cas d'un script
    corrigé (ex. 10_external_baselines.py, MENDEZ ALVARO réintégrée) relancé
    après le run initial fautif. raw_results.csv n'est jamais modifié : ce
    filtre n'agit qu'en mémoire, sur la vue utilisée pour les tables.

    Ne touche à AUCUN autre cas de doublon : si un groupe est 100% SUSPECT
    ou 100% propre malgré plusieurs run_id, il reste tel quel et
    `assert_no_duplicate_conditions_across_run_ids` le signalera comme
    avant — c'est délibéré, cette assertion existe pour attraper les VRAIS
    bugs (incident k=5, cf. CHANGELOG_TABLES.md) et ne doit pas être
    affaiblie par ce mécanisme de supersession volontaire."""
    key_str = df[IDENTITY_COLS].fillna("__NA__").astype(str).agg("||".join, axis=1)
    is_suspect = df["provenance_note"].fillna("").str.contains("|".join(_SUSPECT_MARKERS), regex=True)
    tmp = pd.DataFrame({"key": key_str, "run_id": df["run_id"].values, "suspect": is_suspect.values},
                       index=df.index)
    drop_idx = []
    superseded = []
    for k, g in tmp.groupby("key"):
        if g["run_id"].nunique() <= 1:
            continue
        if g["suspect"].all() or (~g["suspect"]).all():
            continue  # pas un cas de supersession — laissé à l'assertion existante
        bad = g[g["suspect"]]
        drop_idx.extend(bad.index.tolist())
        superseded.append((k, sorted(bad["run_id"].unique()), sorted(g.loc[~g["suspect"], "run_id"].unique())))
    if superseded:
        print(f"  {len(superseded)} condition(s) SUSPECT supersédée(s) par un re-run propre "
             f"(lignes SUSPECT exclues des tables, conservées telles quelles dans raw_results.csv) :")
        for k, old_ids, new_ids in superseded[:10]:
            print(f"    {k} -> run_id(s) exclu(s) : {old_ids} ; retenu(s) : {new_ids}")
    return df.drop(index=drop_idx) if drop_idx else df


# Incident spécifique (P11.6, 2026-08-24) : les 15 reruns du correctif
# NaN-sort (build_correlation_graph, commit 9992793) n'ont persisté QUE
# l'agrégat pour Madrid/correlation/k=5/GCN-Transformer —
# e8_k_sensitivity_3seeds.py (l'outil utilisé pour ce rerun) n'écrit jamais
# de lignes par-station, seulement station="__aggregate__". Table 6
# (Cohen's d/Wilcoxon, calculés par station) et Table 8 continuaient donc de
# lire les lignes par-station de l'ancien run migrated_06_train_multistation_madrid
# (graphe pré-correctif, 24 arêtes au lieu de 30 à k=5) — trouvé en vérifiant
# ce changelog, pas une régression de ce commit-ci. Un rerun dédié
# (06_train_multistation.py --city madrid --graph correlation) fournit des
# lignes par-station fraîches sous le graphe corrigé.
#
# Linear-Transformer est topology-independent (vérifié : ses 7 valeurs par
# station sont identiques entre l'ancien run et ce rerun) — ses lignes ne
# sont pas remplacées, seulement les 21 lignes par-station GCN-Transformer.
# L'agrégat GCN de ce rerun (0.4033±0.0215, backend MPS, 3 seeds) N'EST PAS
# utilisé comme référence Table 4/6/9 : cette place reste occupée par
# l'agrégat déjà établi par le rerun e8 (0.3996±0.0335, backend CPU) pour ne
# pas introduire une seconde valeur candidate sur un chiffre déjà documenté
# dans CHANGELOG_TABLES.md — l'écart entre les deux (0.0037, très inférieur
# aux deux écarts-types) est cohérent avec la non-déterminisme MPS vs CPU
# déjà toléré ailleurs (cf. assert_pruning_anchor_equals_t2, tolérance
# 3×std inter-seeds).
_MADRID_CORR_NAN_BUG_STALE_RUN_ID = "migrated_06_train_multistation_madrid"
_MADRID_CORR_NAN_BUG_FRESH_RUN_ID = "06_madrid_1787591639_ab7abd7b"


def resolve_madrid_correlation_nan_bug_per_station_rows(df):
    stale = ((df.run_id == _MADRID_CORR_NAN_BUG_STALE_RUN_ID) &
            (df.city == "madrid") & (df.topology == "correlation") &
            (df.model == "GCN-Transformer") & (df.station != "__aggregate__"))
    fresh_run = (df.run_id == _MADRID_CORR_NAN_BUG_FRESH_RUN_ID)
    fresh_keep = fresh_run & (df.model == "GCN-Transformer") & (df.station != "__aggregate__")
    fresh_drop = fresh_run & ~fresh_keep
    drop = stale | fresh_drop
    if drop.any():
        print(f"  Madrid/correlation/GCN-Transformer par-station : {int(stale.sum())} ligne(s) "
             f"pré-correctif NaN-sort exclues (run_id={_MADRID_CORR_NAN_BUG_STALE_RUN_ID}), "
             f"remplacées par {int(fresh_keep.sum())} ligne(s) du rerun "
             f"{_MADRID_CORR_NAN_BUG_FRESH_RUN_ID} (agrégat/Linear-Transformer de ce rerun exclus, "
             f"référence Table 4/6/9 inchangée).")
    return df.drop(index=df.index[drop]) if drop.any() else df


def agg_rows(df, model, variant="", topology=None, k=None, keep_frac=None):
    m = (df.station == "__aggregate__") & (df.model == model) & (df.variant == variant)
    if topology is not None:
        m &= (df.topology == topology)
    if k is not None:
        m &= (df.k == k)
    if keep_frac is not None:
        m &= (df.keep_frac == keep_frac)
    return df[m]


def station_rows(df, model, variant="", topology=None, k=None, seed=None):
    m = (df.station != "__aggregate__") & (df.model == model) & (df.variant == variant)
    if topology is not None:
        m &= (df.topology == topology)
    if k is not None:
        m &= (df.k == k)
    if seed is not None:
        m &= (df.seed.astype(str) == str(seed))
    return df[m]


def k_sensitivity_delta(df, city, topology, k):
    """Chemin de calcul INDÉPENDANT de gcn_lin_3seed (dicts par seed, pas de
    passage par pandas .sort_values/zip) — utilisé par Table 9 ET par
    l'assertion T9==T6, pour que la comparaison teste deux implémentations
    réellement différentes de la même quantité, pas la même fonction appelée
    deux fois (cf. rapport de tâche P4)."""
    gcn = agg_rows(df, "GCN-Transformer", topology=topology, k=k)
    gcn = gcn[gcn.city == city]
    lin = agg_rows(df, "Linear-Transformer", topology=topology)
    lin = lin[lin.city == city]
    if len(gcn) == 0 or len(lin) == 0:
        return None
    gcn_seeds = {str(s): r2 for s, r2 in zip(gcn.seed, gcn.r2)}
    lin_seeds = {str(s): r2 for s, r2 in zip(lin.seed, lin.r2)}
    common = sorted(set(gcn_seeds) & set(lin_seeds))
    if not common:
        return None
    delta = np.array([gcn_seeds[s] for s in common]) - np.array([lin_seeds[s] for s in common])
    dm, ds = agg_mean_std(delta)
    suspect = _is_suspect(gcn.provenance_note)
    return dict(delta_mean=dm, delta_std=ds, n_seeds=len(common), suspect=suspect)


def gcn_lin_3seed(df, city, topology):
    """(gcn_mean, gcn_std, lin_mean, lin_std, delta_mean, delta_std) k=5,
    3 seeds, ddof=1 — la quantité centrale de T4/T6/T9(k=5)."""
    gcn = agg_rows(df, "GCN-Transformer", topology=topology, k=5)
    gcn = gcn[gcn.city == city].sort_values("seed")
    lin = agg_rows(df, "Linear-Transformer", topology=topology)
    lin = lin[lin.city == city].sort_values("seed")
    if len(gcn) == 0 or len(lin) == 0:
        return None
    gcn_seeds = {int(float(s)): r2 for s, r2 in zip(gcn.seed, gcn.r2)}
    lin_seeds = {int(float(s)): r2 for s, r2 in zip(lin.seed, lin.r2)}
    common = sorted(set(gcn_seeds) & set(lin_seeds))
    if not common:
        return None
    gcn_vals = np.array([gcn_seeds[s] for s in common])
    lin_vals = np.array([lin_seeds[s] for s in common])
    delta_vals = gcn_vals - lin_vals
    gm, gs = agg_mean_std(gcn_vals)
    lm, ls = agg_mean_std(lin_vals)
    dm, ds = agg_mean_std(delta_vals)
    return dict(gcn_mean=gm, gcn_std=gs, lin_mean=lm, lin_std=ls,
                delta_mean=dm, delta_std=ds, n_seeds=len(common))


# --------------------------------------------------------------------------- #
# T1 — Caractérisation des réseaux (§3, demande R2.7)
# --------------------------------------------------------------------------- #
CHARACTERIZATION_CSV = ROOT / "analysis" / "p9_3_characterization.csv"


def table_characterization():
    if not CHARACTERIZATION_CSV.exists():
        write_missing("table1_characterization",
                      "Table 1 — Caractérisation des réseaux (R2.7)", 4)
        return
    cdf = pd.read_csv(CHARACTERIZATION_CSV)
    cdf["city"] = cdf.city.str.lower()
    rows = []
    for _, r in cdf.sort_values("city").iterrows():
        rows.append([
            CITY_LABEL.get(r.city, r.city.capitalize()),
            r.period,
            int(r.n_stations),
            r.provider,
            r.weather_source,
            f"{r.pm25_mean:.2f}",
            f"{r.pm25_var:.2f}",
            f"{r.train_var_mean:.2f} [{r.train_var_min:.2f}, {r.train_var_max:.2f}]",
            f"{r.lag1_autocorr:.4f}",
            f"{r.r2_persistence:.4f}",
            f"{r.raw_missing_rate:.4f}",
            f"{r.density_distance:.4f}",
            f"{r.density_correlation:.4f}",
            f"{r.degree_eff_distance:.2f}",
            f"{r.degree_eff_correlation:.2f}",
            f"{r.r_bar:.4f}",
        ])
    write_table("table1_characterization",
               "Table 1 — Caractérisation des réseaux (R2.7)",
               ["City", "Period", "Stations", "Provider", "Weather source",
                "PM2.5 mean", "PM2.5 var", "Train var/station [min,max]",
                "Lag-1 autocorr", "R² persistence", "Raw missing rate",
                "Density (distance)", "Density (correlation)",
                "Degree eff. (distance)", "Degree eff. (correlation)", "r̄"],
               rows,
               "Source : analysis/p9_3_characterization.csv (P9.3). Train var/station : "
               "variance PM2.5 par station sur la période train (70% initiaux), "
               "moyenne puis [min, max] inter-stations — Madrid min=0.00 correspond à "
               "MENDEZ ALVARO (PM2.5 constant sur train, cf. REVISION_BRIEF.md). r̄ et "
               "densité/degré effectif calculés sur la période train (70% initiaux), "
               "à ne pas confondre avec h(D) (jeu complet, Table 2).")


# --------------------------------------------------------------------------- #
# T2 — h(D)  [EXCEPTION : source = results/heterogeneity_index_v2.csv]
# --------------------------------------------------------------------------- #
_H_INDEX_CITY_CODE = {"chang-zhu-tan": "czt"}  # label CSV (affichage) -> code city (raw_results.csv)


def table1():
    """Nom de fonction historique — émet désormais Table 2 (h(D)), cf.
    numérotation dans le docstring du module."""
    if not H_INDEX_PATH.exists():
        write_missing("table2_h_index", "Table 2 — Indice d'hétérophilie spatiale h(D)", 4)
        return None
    hdf = pd.read_csv(H_INDEX_PATH)
    hdf["city"] = hdf.city.str.lower().map(lambda c: _H_INDEX_CITY_CODE.get(c, c))
    rows = [[CITY_LABEL.get(r.city, r.city.capitalize()), r.n_stations, f"{r.r_bar:.3f}",
             f"{r.moran_I:.3f}", f"{r.cv_raw:.3f}", f"{r.h:.3f}"]
            for _, r in hdf.sort_values("h").iterrows()]
    write_table("table2_h_index", "Table 2 — Indice d'hétérophilie spatiale h(D)",
               ["City", "Stations", "r̄", "Moran's I", "CV", "h(D)"], rows,
               "Source : results/heterogeneity_index_v2.csv — EXCEPTION documentée, "
               "h(D) n'est pas un résultat de run (hors schéma raw_results.csv, cf. P3).")
    return hdf.set_index("city")["h"].to_dict()


# --------------------------------------------------------------------------- #
# T4 — Benchmark canonique
# --------------------------------------------------------------------------- #
def table2(df):
    rows = []
    for city in CITIES_4NET:
        for topo in TOPOS:
            r = gcn_lin_3seed(df, city, topo)
            if r is None:
                rows.append([CITY_LABEL[city], topo, "MISSING", "MISSING", "MISSING"])
                continue
            rows.append([CITY_LABEL[city], topo,
                        f"{r['lin_mean']:.4f} ± {r['lin_std']:.4f}",
                        f"{r['gcn_mean']:.4f} ± {r['gcn_std']:.4f}",
                        f"{r['delta_mean']:+.4f} ± {r['delta_std']:.4f}"])
    write_table("table4_benchmark", "Table 4 — Benchmark canonique (k=5, 3 seeds, ddof=1)",
               ["City", "Topology", "Linear-Transformer R²", "GCN-Transformer R²", "ΔR²"], rows)


# --------------------------------------------------------------------------- #
# T5 — Baselines externes (4 décimales + colonne provenance)
# --------------------------------------------------------------------------- #
def table3(df):
    rows = []
    for city in CITIES:
        for model, label, protocol in [
            ("Persistence", "Persistence (t−1)", "deterministic"),
            ("ARIMA", "ARIMA", "deterministic"),
            ("XGBoost", "XGBoost", "primary_seed"),
            ("LSTM", "LSTM", "3seed_mean"),
        ]:
            sub = agg_rows(df, model)
            sub = sub[sub.city == city]
            if len(sub) == 0:
                rows.append([city.capitalize(), label, "", "MISSING", "MISSING", "MISSING", protocol])
                continue
            m, s = agg_mean_std(sub.r2.values)
            mae_m, _ = agg_mean_std(sub.mae.values)
            rmse_m, _ = agg_mean_std(sub.rmse.values)
            flag = " [SUSPECT]" if _is_suspect(sub.provenance_note) else ""
            rows.append([city.capitalize(), label, "", f"{mae_m:.4f}", f"{rmse_m:.4f}",
                        f"{m:.4f} ± {s:.4f}{flag}", protocol])
        for topo in TOPOS:
            r = gcn_lin_3seed(df, city, topo)
            if r is None:
                continue
            rows.append([city.capitalize(), "GCN-Transformer", topo, "", "",
                        f"{r['gcn_mean']:.4f} ± {r['gcn_std']:.4f}", "3seed_mean"])
        lin = agg_rows(df, "Linear-Transformer", topology="distance")
        lin = lin[lin.city == city]
        if len(lin):
            m, s = agg_mean_std(lin.r2.values)
            rows.append([city.capitalize(), "Linear-Transformer", "n/a (topology-independent)",
                        "", "", f"{m:.4f} ± {s:.4f}", "3seed_mean"])
        for model, label in [("stgcn", "STGCN"), ("graphwavenet", "Graph WaveNet")]:
            sub = agg_rows(df, model, topology="correlation")
            sub = sub[sub.city == city]
            if len(sub) == 0:
                rows.append([city.capitalize(), label, "correlation", "MISSING", "MISSING",
                            "MISSING", "3seed_mean"])
                continue
            protocol = "3seed_mean" if len(sub) == 3 else "primary_seed"
            if protocol == "3seed_mean":
                m, s = agg_mean_std(sub.r2.values)
                r2_str = f"{m:.4f} ± {s:.4f}"
            else:
                r2_str = f"{sub.r2.iloc[0]:.4f}"
            rows.append([city.capitalize(), label, "correlation", "", "", r2_str, protocol])
    write_table("table5_external_baselines",
               "Table 5 — Baselines externes + SOTA (4 décimales)",
               ["City", "Model", "Topology", "MAE", "RMSE", "R²", "Provenance"], rows,
               "Provenance : 3seed_mean (moyenne±SD ddof=1 sur seeds 42/123/777) ou "
               "primary_seed (seed 42 seul, coût de calcul) ou deterministic (ARIMA/Persistence). "
               "[SUSPECT] : au moins une ligne source porte un provenance_note "
               "(SUSPECT_6STATION, Madrid pré-E1 re-run) — cf. CHANGELOG_TABLES.md.")
    return rows


# --------------------------------------------------------------------------- #
# T6 — Tests statistiques (Wilcoxon, bootstrap CI, Cohen's d, Holm-Bonferroni)
# --------------------------------------------------------------------------- #
def table4(df):
    raw_p = []
    entries = []
    for city in CITIES_4NET:
        for topo in TOPOS:
            gcn_ps = station_rows(df, "GCN-Transformer", topology=topo, k=5, seed=PRIMARY_SEED)
            gcn_ps = gcn_ps[gcn_ps.city == city]
            lin_ps = station_rows(df, "Linear-Transformer", topology=topo, seed=PRIMARY_SEED)
            lin_ps = lin_ps[lin_ps.city == city]
            common = sorted(set(gcn_ps.station) & set(lin_ps.station))
            r3 = gcn_lin_3seed(df, city, topo)
            if not common or r3 is None:
                entries.append(dict(city=city, topology=topo, missing=True))
                continue
            g = gcn_ps.set_index("station").loc[common].r2.values.astype(float)
            l = lin_ps.set_index("station").loc[common].r2.values.astype(float)
            w = stats.wilcoxon(g, l, alternative="less")
            # Cohen's d, ddof=1 (formule standard pour un d apparié) via
            # agg_mean_std — pas un agrégat inter-seeds, mais on route quand
            # même par la fonction unique plutôt qu'un appel nu à l'écart-type
            # (cf. tests/test_ddof.py). NB : results/statistical_analysis/*
            # (pré-P4) utilisait ddof=0 pour ce même calcul (hors périmètre du
            # correctif ddof=1 de P2, qui ne visait que l'agrégation inter-
            # seeds) — cette table régénérée utilise donc délibérément la
            # formule ddof=1 standard, qui peut différer légèrement des
            # valeurs Cohen's d publiées dans le cycle précédent.
            diff = g - l
            diff_mean, diff_std = agg_mean_std(diff)
            d = float(diff_mean / diff_std) if diff_std > 0 else 0.0
            entries.append(dict(city=city, topology=topo, missing=False,
                                delta_mean=r3["delta_mean"], delta_std=r3["delta_std"],
                                pvalue=w.pvalue, cohens_d=d,
                                n_worse=int((g < l).sum()), n_total=len(common)))
            raw_p.append(w.pvalue)
    # Holm-Bonferroni sur les p-values disponibles
    order = np.argsort(raw_p)
    m = len(raw_p)
    corrected = [None] * m
    running_max = 0.0
    for rank, idx in enumerate(order):
        adj = (m - rank) * raw_p[idx]
        running_max = max(running_max, adj)
        corrected[idx] = min(running_max, 1.0)
    pi = 0
    rows = []
    for e in entries:
        if e["missing"]:
            rows.append([CITY_LABEL[e["city"]], e["topology"], "MISSING", "MISSING",
                        "MISSING", "MISSING", "MISSING"])
            continue
        rows.append([CITY_LABEL[e["city"]], e["topology"],
                    f"{e['delta_mean']:+.4f} ± {e['delta_std']:.4f}",
                    f"{corrected[pi]:.3e}", f"{e['cohens_d']:+.2f}",
                    f"{e['n_worse']}/{e['n_total']}", ""])
        pi += 1
    write_table("table6_statistical_tests",
               "Table 6 — Tests statistiques (Wilcoxon, Holm-Bonferroni, Cohen's d)",
               ["City", "Topology", "ΔR² (3 seeds, ddof=1)", "Wilcoxon p (Holm-Bonf)",
                "Cohen's d", "GCN<Linear / total", ""], rows,
               f"Seed primaire {PRIMARY_SEED} pour Wilcoxon/Cohen's d (per-station) ; "
               "3 seeds ddof=1 pour ΔR² agrégé. Holm-Bonferroni sur l'ensemble des tests disponibles.")
    return entries


# --------------------------------------------------------------------------- #
# T7 — Corrélation inter-villes h(D) vs ΔR²  [utilise T2]
# --------------------------------------------------------------------------- #
def table5(df, h_index):
    rows = []
    if h_index is None:
        write_missing("table7_cross_city_correlation",
                      "Table 7 — Corrélation inter-villes h(D) vs ΔR²", 2)
        return
    for topo in TOPOS:
        h_vals, d_vals, cities_ok = [], [], []
        for city in CITIES_4NET:
            r = gcn_lin_3seed(df, city, topo)
            if r is None or city not in h_index:
                continue
            h_vals.append(h_index[city]); d_vals.append(r["delta_mean"]); cities_ok.append(city)
        if len(h_vals) < 3:
            rows.append([topo, "MISSING", "MISSING", f"n={len(h_vals)}"])
            continue
        rho, p = stats.spearmanr(h_vals, d_vals)
        rows.append([topo, f"{rho:+.3f}", f"{p:.4f}", f"n={len(h_vals)} (descriptif, puissance limitée)"])
    write_table("table7_cross_city_correlation",
               "Table 7 — Corrélation inter-villes h(D) vs ΔR² (Spearman)",
               ["Topology", "ρ", "p", "Note"], rows)


# --------------------------------------------------------------------------- #
# T8 — ΔR² par station
# --------------------------------------------------------------------------- #
def table6(df):
    rows = []
    for city in CITIES_4NET:
        for topo in TOPOS:
            gcn_ps = station_rows(df, "GCN-Transformer", topology=topo, k=5, seed=PRIMARY_SEED)
            gcn_ps = gcn_ps[gcn_ps.city == city]
            lin_ps = station_rows(df, "Linear-Transformer", topology=topo, seed=PRIMARY_SEED)
            lin_ps = lin_ps[lin_ps.city == city]
            common = sorted(set(gcn_ps.station) & set(lin_ps.station))
            if not common:
                rows.append([CITY_LABEL[city], topo, "MISSING", "", "", ""])
                continue
            g = gcn_ps.set_index("station").loc[common].r2.astype(float)
            l = lin_ps.set_index("station").loc[common].r2.astype(float)
            for st in common:
                rows.append([CITY_LABEL[city], topo, st, f"{g[st]:.4f}", f"{l[st]:.4f}",
                            f"{g[st]-l[st]:+.4f}"])
    write_table("table8_per_station", "Table 8 — ΔR² par station (seed primaire 42)",
               ["City", "Topology", "Station", "GCN R²", "Linear R²", "ΔR²"], rows)


# --------------------------------------------------------------------------- #
# T9 — Sensibilité k
# --------------------------------------------------------------------------- #
def table7(df):
    rows = []
    for city in CITIES_4NET:
        for topo in TOPOS:
            cells = []
            for k in KS:
                r = k_sensitivity_delta(df, city, topo, k)
                if r is None:
                    gcn = agg_rows(df, "GCN-Transformer", topology=topo, k=k)
                    gcn = gcn[gcn.city == city]
                    if len(gcn) == 1 and str(gcn.seed.iloc[0]) == "unknown_3seed_agg":
                        cells.append(f"{gcn.r2.iloc[0]:.4f} [UNRECOVERABLE, ddof indisponible]")
                    else:
                        cells.append("MISSING")
                    continue
                flag = " [SUSPECT]" if r["suspect"] else ""
                cells.append(f"{r['delta_mean']:+.4f}±{r['delta_std']:.4f}{flag}")
            rows.append([CITY_LABEL[city], topo] + cells)
    write_table("table9_k_sensitivity", "Table 9 — Sensibilité k (ΔR², 3 seeds, ddof=1)",
               ["City", "Topology", "k=3", "k=5", "k=8"], rows,
               "[SUSPECT] : au moins une ligne source porte un provenance_note "
               "(SUSPECT_6STATION/UNRECOVERABLE) — cf. CHANGELOG_TABLES.md. "
               "CZT : k=5 uniquement (benchmark canonique) — aucun balayage k-sensitivity "
               "lancé sur ce réseau (E16 : protocole limité à GCN-Transformer + "
               "Linear-Transformer), k=3/k=8 MISSING par construction.")


# --------------------------------------------------------------------------- #
# T10 — Over-smoothing / GAT  [MISSING DATA attendu]
# --------------------------------------------------------------------------- #
def table8(df):
    specs = [("Linear-Transformer", "1layer"), ("GCN-Transformer", "1layer"),
             ("GCN-Transformer", "2layer"), ("GAT-Transformer", "2layer")]
    n_present = sum(len(agg_rows(df, m, variant=v, topology="distance")[
                        agg_rows(df, m, variant=v, topology="distance").city == city])
                    for city in CITIES for m, v in specs)
    if n_present == 0:
        write_missing("table10_oversmoothing", "Table 10 — Over-smoothing / GAT", 36)
        return
    rows = []
    for city in CITIES:
        cells = []
        for model, variant in specs:
            sub = agg_rows(df, model, variant=variant, topology="distance")
            sub = sub[sub.city == city]
            if len(sub) == 0:
                cells.append("MISSING")
                continue
            m, s = agg_mean_std(sub.r2.values)
            cells.append(f"{m:.4f}±{s:.4f}")
        rows.append([CITY_LABEL[city]] + cells)
    write_table("table10_oversmoothing", "Table 10 — Over-smoothing / GAT",
               ["City", "Linear (1L)", "GCN (1L)", "GCN (2L)", "GAT (2L)"], rows,
               "Beijing/London/Madrid uniquement (E13) — Chang-Zhu-Tan (CZT) n'a pas de "
               "contrôle over-smoothing/GAT : jamais lancé sur ce réseau (E16 : protocole "
               "limité à GCN-Transformer + Linear-Transformer), pas une omission. "
               "Non iso-capacité avec le GCN-Transformer canonique (Table 4) : cette "
               "réimplémentation (09_controls_oversmoothing.py) utilise un FFN Transformer "
               "de largeur 128 au lieu de 256, soit 35,6 % de paramètres en moins (dont 89 % "
               "imputables au FFN, pas à la profondeur du GCN testée ici) — comparaisons "
               "internes à cette table valides, comparaison directe avec Table 4 confondue "
               "par cet écart. Cf. CHANGELOG_TABLES.md pour le détail chiffré.")


# --------------------------------------------------------------------------- #
# T11 — Contrôles diagnostiques
# --------------------------------------------------------------------------- #
def table9(df):
    rows = []
    for city in CITIES:
        for topo in TOPOS:
            cells = []
            for exp in ["real", "shuffled_graph", "no_meteorology"]:
                gcn = agg_rows(df, "GCN-Transformer", variant=exp, topology=topo)
                gcn = gcn[gcn.city == city]
                lin_variant = "no_meteorology" if exp == "no_meteorology" else ""
                lin = agg_rows(df, "Linear-Transformer", variant=lin_variant,
                              topology=(topo if lin_variant else ""))
                lin = lin[lin.city == city]
                if lin_variant == "":
                    lin = agg_rows(df, "Linear-Transformer", variant="")
                    lin = lin[(lin.city == city) & (lin.topology == topo)]
                if len(gcn) == 0 or len(lin) == 0:
                    cells.append("MISSING")
                    continue
                d = float(gcn.r2.iloc[0]) - float(lin.r2.mean())
                cells.append(f"{d:+.4f}")
            rows.append([CITY_LABEL[city], topo] + cells)
    write_table("table11_diagnostics", "Table 11 — Contrôles diagnostiques (E4/E5, seed 42)",
               ["City", "Topology", "ΔR² real", "ΔR² shuffled_graph", "ΔR² no_meteorology"], rows,
               "Beijing/London/Madrid uniquement (E4/E5) — Chang-Zhu-Tan (CZT) n'a pas de "
               "contrôle diagnostique (shuffled-graph/no-meteorology) : jamais lancé sur ce "
               "réseau (E16 : protocole limité à GCN-Transformer + Linear-Transformer), "
               "pas une omission.")


# --------------------------------------------------------------------------- #
# ASSERTIONS BLOQUANTES
# --------------------------------------------------------------------------- #
def assert_t7_k5_equals_t4(df):
    """Nom de fonction historique — vérifie désormais T9(k=5) == T6.
    T6 (delta via gcn_lin_3seed : pandas sort_values/zip) et T9 (delta via
    k_sensitivity_delta : dicts par seed) sont deux implémentations écrites
    séparément de la même quantité (ΔR² GCN k=5 vs Linear, 3 seeds, ddof=1) —
    une divergence ici indique un vrai bug (mauvais filtrage de seed, mauvais
    ddof, désalignement), pas une tautologie de code partagé."""
    for city in CITIES_4NET:
        for topo in TOPOS:
            name = f"T9(k=5) == T6 [{city}/{topo}]"
            t4 = gcn_lin_3seed(df, city, topo)
            t7 = k_sensitivity_delta(df, city, topo, 5)
            if t4 is None or t7 is None:
                record(name, "MISSING DATA",
                      f"T6 présent={t4 is not None}, T9(k=5) présent={t7 is not None}")
                continue
            mean_diff = abs(t4["delta_mean"] - t7["delta_mean"])
            std_diff = abs(t4["delta_std"] - t7["delta_std"])
            ok = mean_diff < TOL_EXACT and std_diff < TOL_EXACT
            record(name, "PASS" if ok else "FAIL",
                  f"mean Δ={mean_diff:.2e} (tol {TOL_EXACT:.0e}), std Δ={std_diff:.2e}, "
                  f"T6={t4['delta_mean']:.6f}±{t4['delta_std']:.6f}, "
                  f"T9={t7['delta_mean']:.6f}±{t7['delta_std']:.6f}")


def assert_pruning_anchor_equals_t2(df):
    for city in CITIES:
        for seed in SEEDS:
            name = f"pruning anchor (keep_frac=1.0) == T4 GCN/distance [{city}/seed={seed}]"
            anchor = agg_rows(df, "GCN-Transformer", topology="distance", keep_frac=1.0)
            anchor = anchor[(anchor.city == city) & (anchor.seed.astype(str) == str(seed))]
            t2 = agg_rows(df, "GCN-Transformer", topology="distance", k=5)
            t2 = t2[(t2.city == city) & (t2.seed.astype(str) == str(seed))]
            if len(anchor) == 0 or len(t2) == 0:
                record(name, "MISSING DATA",
                      f"anchor présent={len(anchor)>0}, T4 présent={len(t2)>0}")
                continue
            a_val = float(anchor.r2.iloc[0]); t2_val = float(t2.r2.iloc[0])
            diff = abs(a_val - t2_val)
            # Tolérance = 3x l'écart-type inter-seeds observé pour cette ville/
            # topologie (bruit normal d'entraînement, PAS ajusté après coup :
            # dérivé de T9 k=5, jamais de cette comparaison elle-même) — ce
            # sont deux entraînements INDÉPENDANTS du même graphe, pas la même
            # donnée relue deux fois (contrairement à l'assertion précédente).
            r3 = gcn_lin_3seed(df, city, "distance")
            noise_scale = r3["gcn_std"] if r3 else 0.0
            tol = max(3 * noise_scale, 0.01)
            ok = diff < tol
            anchor_n, t2_n = int(anchor.n_stations.iloc[0]), int(t2.n_stations.iloc[0])
            cause = ""
            if not ok and anchor_n != t2_n:
                cause = (f" — CAUSE IDENTIFIÉE : n_stations diffère ({anchor_n} vs {t2_n}), "
                        "cohérent avec le SUSPECT_6STATION déjà documenté en P1 (MENDEZ ALVARO "
                        "absente du pruning Madrid, ancien protocole) — pas un nouveau bug, "
                        "attend le re-run E10 (cf. CHANGELOG_TABLES.md).")
            record(name, "PASS" if ok else "FAIL",
                  f"|Δ|={diff:.4f}, tol={tol:.4f} (3×std inter-seeds GCN/distance={noise_scale:.4f}), "
                  f"anchor={a_val:.4f} (n={anchor_n}), T4={t2_val:.4f} (n={t2_n}){cause}")


def assert_single_linear_reference(df, t3_rows, t4_entries):
    for city in CITIES_4NET:
        for topo in TOPOS:
            name = f"référence Linear-Transformer unique [{city}/{topo}, 3seed_mean]"
            lin = agg_rows(df, "Linear-Transformer", topology=topo)
            lin = lin[lin.city == city]
            if len(lin) == 0:
                record(name, "MISSING DATA", "aucune ligne Linear-Transformer")
                continue
            distinct = lin.groupby("seed").r2.first()
            m, _ = agg_mean_std(distinct.values)
            # cohérence avec la valeur utilisée dans T6 pour cette même ville/topo
            entry = next((e for e in t4_entries if e["city"] == city and e["topology"] == topo
                         and not e["missing"]), None)
            if entry is None:
                record(name, "MISSING DATA", "T6 n'a pas cette condition")
                continue
            r3 = gcn_lin_3seed(df, city, topo)
            diff = abs(r3["lin_mean"] - m)
            record(name, "PASS" if diff < TOL_EXACT else "FAIL", f"|Δ|={diff:.2e}")


def assert_t3_graph_models_declare_topology(t3_rows):
    name = "T5 : toute entrée de modèle graphe déclare sa topologie"
    graph_labels = {"GCN-Transformer", "STGCN", "Graph WaveNet"}
    offenders = [r for r in t3_rows if r[1] in graph_labels and (not r[2] or r[2] == "MISSING")]
    if not t3_rows:
        record(name, "MISSING DATA", "Table 5 vide")
    elif offenders:
        record(name, "FAIL", f"{len(offenders)} ligne(s) sans topologie : {offenders[:3]}")
    else:
        record(name, "PASS", f"{sum(1 for r in t3_rows if r[1] in graph_labels)} lignes vérifiées")


IDENTITY_COLS = ["city", "model", "topology", "k", "keep_frac", "variant", "seed", "station", "split"]


def assert_no_duplicate_conditions_across_run_ids(df):
    """Aucune condition logique (IDENTITY_COLS) ne doit apparaître sous plus
    d'un run_id. C'est le trou exact par lequel le doublon k=5 est passé :
    append_run refuse une clé déjà présente, mais sa clé inclut run_id — deux
    run_id différents pour la MÊME condition ne sont jamais détectés comme
    collision. Cette assertion vérifie l'identité logique, pas la clé
    d'écriture."""
    name = "unicité structurelle : une condition = un seul run_id"
    # Comparaison en str (fillna d'abord) : évite tout piège de dtype (NaN vs
    # None vs "", float vs object) au moment du groupby ET du ré-examen des
    # lignes en cas de doublon — un seul et même dataframe stringifié utilisé
    # de bout en bout, jamais de merge sur les dtypes bruts d'origine.
    key = df[IDENTITY_COLS].fillna("__NA__").astype(str)
    tmp = key.copy()
    tmp["run_id"] = df["run_id"].values
    grouped = tmp.groupby(IDENTITY_COLS, dropna=False)["run_id"].nunique()
    dup_keys = grouped[grouped > 1]
    if len(dup_keys) == 0:
        record(name, "PASS", f"{len(df)} lignes, {len(grouped)} conditions distinctes, 0 doublon")
        return
    examples = []
    for idx in dup_keys.index[:10]:
        idx_t = idx if isinstance(idx, tuple) else (idx,)
        mask = pd.Series(True, index=tmp.index)
        for col, val in zip(IDENTITY_COLS, idx_t):
            mask &= (tmp[col] == val)
        run_ids = sorted(tmp.loc[mask, "run_id"].unique())
        examples.append(f"{dict(zip(IDENTITY_COLS, idx_t))} -> run_ids={run_ids}")
    record(name, "FAIL",
          f"{len(dup_keys)} condition(s) avec >1 run_id : " + " | ".join(examples))


def assert_effective_degree(df):
    name = "degré effectif = min(k, N-1) ; Beijing jamais plafonné pour k≤11"
    sub = df[(df.city == "beijing") & df.k.notna() & (df.model == "GCN-Transformer")]
    if len(sub) == 0:
        record(name, "MISSING DATA", "aucune ligne Beijing avec k")
        return
    n_stations = sub.n_stations.dropna().unique()
    if len(n_stations) == 0:
        record(name, "MISSING DATA", "n_stations absent pour Beijing")
        return
    n = int(n_stations[0])
    capped = sub[sub.k > n - 1]
    ok = (len(capped) == 0) and all(k <= 11 for k in sub.k.unique())
    record(name, "PASS" if ok else "FAIL",
          f"n_stations={n}, k_values={sorted(sub.k.unique())}, "
          f"{len(capped)} ligne(s) plafonnée(s)")


def assert_unrounded_recomputation(df):
    """Auto-test à deux volets :
    (1) démontre que ça compte : sur un exemple réel, le delta recalculé
        depuis les R² bruts diffère du delta recalculé depuis les mêmes R²
        pré-arrondis à 4 décimales (sinon le test ne prouverait rien) ;
    (2) vérifie que gcn_lin_3seed/k_sensitivity_delta (utilisés par TOUTES
        les tables et assertions de ce script) produisent bien la version
        brute, pas la version arrondie."""
    name = "différences recalculées depuis les valeurs non arrondies (pas depuis l'affichage)"
    gcn = agg_rows(df, "GCN-Transformer", topology="distance", k=5)
    gcn = gcn[gcn.city == "madrid"].sort_values("seed").reset_index(drop=True)
    lin = agg_rows(df, "Linear-Transformer", topology="distance")
    lin = lin[lin.city == "madrid"].sort_values("seed").reset_index(drop=True)
    if len(gcn) == 0 or len(lin) == 0:
        record(name, "MISSING DATA", "madrid/distance k=5 absent")
        return
    raw_delta = float(gcn.r2.iloc[0]) - float(lin.r2.iloc[0])
    rounded_delta = round(float(gcn.r2.iloc[0]), 4) - round(float(lin.r2.iloc[0]), 4)
    test_is_meaningful = abs(raw_delta - rounded_delta) > 1e-9
    r = gcn_lin_3seed(df, "madrid", "distance")
    table_uses_raw = abs(r["gcn_mean"] - float(gcn.r2.mean())) < 1e-12
    ok = test_is_meaningful and table_uses_raw
    record(name, "PASS" if ok else "FAIL",
          f"exemple madrid/distance/seed42 : brut={raw_delta:.6f} vs arrondi-4dp={rounded_delta:.6f} "
          f"(écart {abs(raw_delta-rounded_delta):.2e}, {'significatif' if test_is_meaningful else 'NUL — test non probant'}); "
          f"gcn_lin_3seed utilise les valeurs brutes : {table_uses_raw}")


def main():
    print("Chargement raw_results.csv ...")
    df = load()
    print(f"{len(df)} lignes.\n")

    print("=== Génération des tables ===")
    table_characterization()
    h_index = table1()
    table2(df)
    t3_rows = table3(df)
    t4_entries = table4(df)
    table5(df, h_index)
    table6(df)
    table7(df)
    table8(df)
    table9(df)

    print("\n=== Assertions bloquantes ===")
    assert_no_duplicate_conditions_across_run_ids(df)
    assert_t7_k5_equals_t4(df)
    assert_pruning_anchor_equals_t2(df)
    assert_single_linear_reference(df, t3_rows, t4_entries)
    assert_t3_graph_models_declare_topology(t3_rows)
    assert_effective_degree(df)
    assert_unrounded_recomputation(df)

    n_pass = sum(1 for a in ASSERTIONS if a["status"] == "PASS")
    n_fail = sum(1 for a in ASSERTIONS if a["status"] == "FAIL")
    n_miss = sum(1 for a in ASSERTIONS if a["status"] == "MISSING DATA")
    print(f"\n{'='*60}\nRÉCAP ASSERTIONS : {n_pass} PASS, {n_fail} FAIL, {n_miss} MISSING DATA "
          f"(sur {len(ASSERTIONS)})\n{'='*60}")
    for a in ASSERTIONS:
        print(f"  [{a['status']:>13}] {a['name']}")

    if n_fail:
        print(f"\n{n_fail} assertion(s) FAIL — corriger les données/le pipeline, pas cette liste.")
        sys.exit(1)


if __name__ == "__main__":
    main()
